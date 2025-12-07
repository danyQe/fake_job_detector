from fastapi import APIRouter, Depends, HTTPException, Body, UploadFile, File, Query, status
from sqlalchemy.orm import Session
from datetime import datetime
import google.genai as genai
from google.genai import types
import os
from app.services.resume_service import ResumeService
import json
from dotenv import load_dotenv
from typing import List, Optional
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile
from fastapi.responses import FileResponse
import weasyprint
from app.services.pdf_parser import PDFParser
import PyPDF2
import io

from app import schemas, models, auth
from app.database import get_db

load_dotenv()

# Get environment variables
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
# genai.configure(api_key=GEMINI_API_KEY)

router = APIRouter(
    prefix="/resumes",
    tags=["Resume Generation"],
    responses={404: {"description": "Not found"}},
)

# Initialize PDF parser
pdf_parser = PDFParser()

# LLM prompt for resume generation
RESUME_GENERATION_PROMPT = """
You are a professional resume writer with expertise in creating ATS-optimized resumes. 
Based on the provided user information and job details, create an enhanced, tailored resume.

Job Details:
{job_details}

User's Current Information:
{personal_info}

Your task is to enhance this resume to better match the job requirements while maintaining truthfulness.
For each work experience, provide bullet points that highlight relevant achievements and responsibilities.
Enhance the professional summary to be more compelling and targeted to this specific role.
Add relevant skills that match the job requirements based on the user's background.

If the user is a fresher or has no work experience entries:
1. Focus on creating a stronger professional summary highlighting education, skills, and career objectives
2. Emphasize relevant academic projects, coursework, and achievements
3. Highlight internships or volunteer work if available
4. Suggest transferable skills from academic or extracurricular activities that match the job requirements
5. Leave the experience array empty or include only relevant internships/projects

Format your response as a structured JSON object following this schema:
{{
  "fullName": "User's full name",
  "title": "Professional title tailored to the job",
  "email": "User's email",
  "phone": "User's phone number",
  "location": "City, State",
  "linkedin": "LinkedIn URL (optional)",
  "summary": "Enhanced professional summary tailored to the job (3-5 sentences)",
  "skills": ["Skill 1", "Skill 2", "Skill 3", ...],
  "experience": [
    {{
      "company": "Company name",
      "title": "Job title",
      "startDate": "MM/YYYY",
      "endDate": "MM/YYYY or null if current",
      "current": true/false,
      "description": "Job responsibilities and achievements in bullet points format"
    }},
    ...
  ],
  "education": [
    {{
      "institution": "School name",
      "degree": "Degree name",
      "field": "Field of study",
      "graduationDate": "MM/YYYY"
    }},
    ...
  ],
  "certifications": [
    {{
      "name": "Certification name",
      "issuer": "Issuing organization"
    }},
    ...
  ]
}}

Ensure the JSON is valid and properly formatted.
"""


def generate_docx_resume(resume_data):
    doc = docx.Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Name and Title
    name = doc.add_paragraph()
    name_run = name.add_run(resume_data["fullName"])
    name_run.bold = True
    name_run.font.size = Pt(18)
    name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    title = doc.add_paragraph()
    title_run = title.add_run(resume_data["title"])
    title_run.font.size = Pt(14)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Contact Info
    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact.add_run(f"{resume_data['email']} | {resume_data['phone']} | {resume_data['location']}")
    if resume_data.get("linkedin"):
        contact.add_run(f" | {resume_data['linkedin']}")
    
    doc.add_paragraph()
    
    # Summary
    summary_heading = doc.add_heading("PROFESSIONAL SUMMARY", level=1)
    summary_heading.style.font.size = Pt(12)
    summary_heading.style.font.bold = True
    doc.add_paragraph(resume_data["summary"])
    
    # Skills
    skills_heading = doc.add_heading("SKILLS", level=1)
    skills_heading.style.font.size = Pt(12)
    skills_heading.style.font.bold = True
    skills_para = doc.add_paragraph()
    skills_para.add_run(" | ".join(resume_data["skills"]))
    
    # Experience
    exp_heading = doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
    exp_heading.style.font.size = Pt(12)
    exp_heading.style.font.bold = True
    
    for exp in resume_data["experience"]:
        company_para = doc.add_paragraph()
        company_run = company_para.add_run(f"{exp['company']} - {exp['title']}")
        company_run.bold = True
        
        date_para = doc.add_paragraph()
        date_text = f"{exp['startDate']} - "
        date_text += "Present" if exp.get("current") else exp['endDate']
        date_para.add_run(date_text)
        
        # Split bullet points by newline or assume it's a paragraph
        desc_text = exp["description"]
        if "\n" in desc_text:
            bullets = desc_text.split("\n")
            for bullet in bullets:
                if bullet.strip():
                    bullet_para = doc.add_paragraph(bullet.strip(), style='List Bullet')
        else:
            doc.add_paragraph(desc_text)
    
    # Education
    edu_heading = doc.add_heading("EDUCATION", level=1)
    edu_heading.style.font.size = Pt(12)
    edu_heading.style.font.bold = True
    
    for edu in resume_data["education"]:
        edu_para = doc.add_paragraph()
        edu_run = edu_para.add_run(f"{edu['institution']} - {edu['degree']} in {edu['field']}")
        edu_run.bold = True
        
        grad_para = doc.add_paragraph()
        grad_para.add_run(f"Graduated: {edu['graduationDate']}")
    
    # Certifications (if any)
    if resume_data.get("certifications"):
        cert_heading = doc.add_heading("CERTIFICATIONS", level=1)
        cert_heading.style.font.size = Pt(12)
        cert_heading.style.font.bold = True
        
        for cert in resume_data["certifications"]:
            cert_para = doc.add_paragraph()
            cert_para.add_run(f"{cert['name']} - {cert['issuer']}")
    
    # Save document to temp file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name

def generate_pdf_resume(resume_data):
    html_template = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: 'Arial', sans-serif;
                font-size: 11pt;
                margin: 0.7in;
                color: #000;
            }}
            .header {{
                text-align: center;
                margin-bottom: 10px;
            }}
            .header h1 {{
                margin: 0;
                font-size: 16pt;
                font-weight: bold;
            }}
            .contact {{
                font-size: 10pt;
                margin-top: 5px;
                line-height: 1.4;
            }}
            h2 {{
                background-color: #f0f0f0;
                padding: 4px 8px;
                font-size: 11pt;
                margin-top: 20px;
                margin-bottom: 8px;
            }}
            .section p {{
                margin: 3px 0;
            }}
            .item-title {{
                font-weight: bold;
                font-size: 11pt;
            }}
            .item-sub {{
                font-style: italic;
                font-size: 10pt;
                margin-bottom: 2px;
            }}
            ul {{
                margin-top: 5px;
                margin-bottom: 10px;
                padding-left: 20px;
            }}
            .two-col {{
                display: flex;
                justify-content: space-between;
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>{resume_data["fullName"]}</h1>
            <div class="contact">
                {resume_data["email"]} | {resume_data["phone"]} | {resume_data["location"]}<br>
                {resume_data.get("linkedin", "")} {resume_data.get("portfolio", "")} {resume_data.get("github", "")}
            </div>
        </div>

        <h2>OBJECTIVE</h2>
        <p>{resume_data["summary"]}</p>

        <h2>EDUCATION</h2>
    """
    for edu in resume_data["education"]:
        html_template += f"""
            <div class="two-col">
                <div class="item-title">{edu["institution"]}</div>
                <div>{edu["graduationDate"]}</div>
            </div>
            <p class="item-sub">{edu["degree"]} in {edu["field"]}</p>
        """

    html_template += "<h2>EXPERIENCE</h2>"
    for exp in resume_data["experience"]:
        date_range = f"{exp['startDate']} – {'Present' if exp.get('current') else exp['endDate']}"
        html_template += f"""
            <div class="two-col">
                <div class="item-title">{exp["company"]} - {exp["title"]}</div>
                <div>{date_range}</div>
            </div>
        """
        bullets = exp["description"].split("\n")
        html_template += "<ul>"
        for bullet in bullets:
            if bullet.strip():
                html_template += f"<li>{bullet.strip()}</li>"
        html_template += "</ul>"

    html_template += "<h2>SKILLS</h2><ul>"
    for skill in resume_data["skills"]:
        html_template += f"<li>{skill}</li>"
    html_template += "</ul>"

    if resume_data.get("projects"):
        html_template += "<h2>PROJECTS</h2>"
        for proj in resume_data["projects"]:
            html_template += f"""
                <p class="item-title">{proj['name']}</p>
                <p class="item-sub">{proj.get('tech', '')}</p>
            """
            bullets = proj["description"].split("\n")
            html_template += "<ul>"
            for bullet in bullets:
                if bullet.strip():
                    html_template += f"<li>{bullet.strip()}</li>"
            html_template += "</ul>"

    if resume_data.get("certifications"):
        html_template += "<h2>CERTIFICATIONS</h2><ul>"
        for cert in resume_data["certifications"]:
            html_template += f"<li>{cert['name']} - {cert['issuer']}</li>"
        html_template += "</ul>"

    html_template += "</body></html>"

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf = weasyprint.HTML(string=html_template).write_pdf()

    with open(temp_file.name, "wb") as f:
        f.write(pdf)

    return temp_file.name

@router.post("/parse-pdf")
async def parse_pdf(
    file: UploadFile = File(...),
    current_user: models.User = Depends(auth.get_current_active_user)
):
    """Parse a PDF resume and return structured data"""
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    
    try:
        # Read the uploaded file
        content = await file.read()
        
        # Extract text from PDF using PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        
        # Use LLM to parse the resume text
        client = genai.Client(api_key=GEMINI_API_KEY)
        
        prompt = f"""
        You are a professional resume parser. Extract structured information from the following resume text.
        Format your response as a JSON object with the following structure:
        {{
          "fullName": "User's full name",
          "title": "Professional title",
          "email": "User's email",
          "phone": "User's phone number",
          "location": "City, State",
          "linkedin": "LinkedIn URL (if available)",
          "summary": "Professional summary (3-5 sentences)",
          "skills": ["Skill 1", "Skill 2", "Skill 3", ...],
          "experience": [
            {{
              "company": "Company name",
              "title": "Job title",
              "startDate": "MM/YYYY",
              "endDate": "MM/YYYY or null if current",
              "current": true/false,
              "description": "Job responsibilities and achievements in bullet points format"
            }},
            ...
          ],
          "education": [
            {{
              "institution": "School name",
              "degree": "Degree name",
              "field": "Field of study",
              "graduationDate": "MM/YYYY"
            }},
            ...
          ],
          "certifications": [
            {{
              "name": "Certification name",
              "issuer": "Issuing organization"
            }},
            ...
          ],
          "projects": [
            {{
              "name": "Project name",
              "technologies": ["Tech 1", "Tech 2", ...],
              "startDate": "MM/YYYY",
              "endDate": "MM/YYYY or null if current",
              "current": true/false,
              "description": "Project description with bullet points",
              "url": "Project URL (optional)"
            }},
            ...
          ]
        }}
        
        If the resume doesn't contain work experience entries (suggesting a fresher profile), set a flag "isFresher": true.
        If the resume clearly mentions the person is a fresher, recent graduate, or doesn't show any work experience, set "isFresher": true.
        
        Ensure that:
        1. Any 'projects' mentioned should be placed under 'projects', not 'experience'
        2. Academic projects should be in 'projects' section, not in 'experience'
        3. Personal projects and GitHub/portfolio projects should be in 'projects' section
        4. If the person is a fresher, set "isFresher": true at the root level of the JSON
        
        Resume text:
        {text}
        """
        
        response = client.models.generate_content(model="gemini-2.0-flash", contents=[prompt])
        resume_text = response.text
        
        # Extract the JSON part from the response
        if "```json" in resume_text:
            json_text = resume_text.split("```json")[1].split("```")[0].strip()
        elif "```" in resume_text:
            json_text = resume_text.split("```")[1].strip()
        else:
            json_text = resume_text.strip()
        
        # Parse the generated resume data
        parsed_data = json.loads(json_text)
        
        # Create empty projects array if none exists
        if "projects" not in parsed_data:
            parsed_data["projects"] = []
            
        # Create empty experience array if none exists
        if "experience" not in parsed_data:
            parsed_data["experience"] = []
            
        # Check if the profile is a fresher
        if "isFresher" not in parsed_data:
            # If experience array is empty or has no valid entries, mark as fresher
            is_empty_experience = (len(parsed_data["experience"]) == 0)
            parsed_data["isFresher"] = is_empty_experience
        
        return parsed_data
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error parsing PDF: {str(e)}")


@router.post("/generate", response_model=schemas.Resume)
async def generate_resume(
    job_id: int = Query(..., description="Job analysis ID to tailor resume for"),
    resume_request: schemas.ResumeRequest = Body(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_active_user)
):
    """Generate a resume and save it to the database"""
    
    # Verify job analysis exists
    job_analysis = db.query(models.JobAnalysis).filter(
        models.JobAnalysis.id == job_id,
        models.JobAnalysis.user_id == current_user.id
    ).first()
    
    if not job_analysis:
        raise HTTPException(status_code=404, detail="Job analysis not found")
    
    try:
        # Generate resume content using LLM
        client = genai.Client(api_key=GEMINI_API_KEY)
        
        # If personal_info is a string (parsed PDF data), parse it
        if isinstance(resume_request.personal_info, str):
            try:
                resume_request.personal_info = json.loads(resume_request.personal_info)
            except json.JSONDecodeError:
                raise HTTPException(status_code=400, detail="Invalid personal info format")
        
        # Check if user is a fresher based on form submission or empty experience
        is_fresher = resume_request.personal_info.get("isFresher", False)
        
        # If user indicated they're a fresher, ensure experience array is empty
        if is_fresher:
            resume_request.personal_info["experience"] = []
        
        # Get job description from job analysis
        job_details = job_analysis.job_content
        
        prompt = RESUME_GENERATION_PROMPT.format(
            job_details=job_details,
            personal_info=json.dumps(resume_request.personal_info)
        )
        print("prompt:",prompt)
        response = client.models.generate_content(model="gemini-2.0-flash", contents=[prompt],config=types.GenerateContentConfig(temperature=0.1))
        resume_text = response.text
        print("resume:",resume_text)
        # Extract the JSON part from the response
        if "```json" in resume_text:
            json_text = resume_text.split("```json")[1].split("```")[0].strip()
        elif "```" in resume_text:
            json_text = resume_text.split("```")[1].strip()
        else:
            json_text = resume_text.strip()
        
        # Parse the generated resume data
        resume_data = json.loads(json_text)
        
        # Add the template from the request to the resume data
        resume_data["template"] = resume_request.template
        
        # Store resume in database
        db_resume = models.Resume(
            user_id=current_user.id,
            job_analysis_id=job_id,
            file_path=None,
            format=resume_request.format
        )
        db.add(db_resume)
        db.commit()
        db.refresh(db_resume)
        
        # Save resume data to a JSON file that can be used later
        resume_data_dir = os.path.join("app", "data", "resumes")
        os.makedirs(resume_data_dir, exist_ok=True)
        resume_data_path = os.path.join(resume_data_dir, f"{db_resume.id}.json")
        
        with open(resume_data_path, "w") as f:
            json.dump(resume_data, f)
        
        # Return with job_analysis_id from the database
        return schemas.Resume(
            id=db_resume.id,
            user_id=db_resume.user_id,
            job_analysis_id=db_resume.job_analysis_id,
            file_path=db_resume.file_path,
            format=db_resume.format,
            created_at=db_resume.created_at
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating resume: {str(e)}")


@router.get("/download/{resume_id}")
async def download_resume(
    resume_id: str,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_active_user)
):
    """Download a generated resume"""
    print(f"\n=== Starting Resume Download Process ===")
    print(f"Resume ID: {resume_id}")
    print(f"User ID: {current_user.id}")
    
    try:
        # Get resume from database
        print("\nFetching resume from database...")
        resume = db.query(*models.Resume.get_compatible_columns()).filter(
            models.Resume.id == resume_id,
            models.Resume.user_id == current_user.id
        ).first()
        
        if not resume:
            print("Resume not found in database")
            raise HTTPException(status_code=404, detail="Resume not found")
        
        print("Resume found in database")
        print(f"Current format: {resume.format}")
        
        try:
            print("\nInitializing resume generator...")
            resume_generator = ResumeService(db)
            
            # Load resume data from JSON file
            resume_data_path = os.path.join("app", "data", "resumes", f"{resume.id}.json")
            print(f"Looking for resume data at: {resume_data_path}")
            
            if not os.path.exists(resume_data_path):
                print("Resume data file not found")
                raise HTTPException(status_code=404, detail="Resume data not found")
            
            print("Loading resume data from JSON file...")
            with open(resume_data_path, "r") as f:
                resume_data = json.load(f)
            print("Resume data loaded successfully")
            
            # Get template from request or use 'modern' as default
            template = resume_data.get("template", "modern")
            print(f"Using template: {template}")
            
            # Generate the file based on format
            print(f"\nGenerating resume in format: {resume.format}")
            file_path = None
            
            try:
                if resume.format.lower() == "docx":
                    print("Generating DOCX format file...")
                    file_path = resume_generator.create_resume_from_data(resume_data, format="docx", template=template)
                    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    filename = f"resume_{resume.id}.docx"
                else:  # PDF format
                    print("Generating PDF format file...")
                    file_path = resume_generator.create_resume_from_data(resume_data, format="pdf", template=template)
                    media_type = "application/pdf"
                    filename = f"resume_{resume.id}.pdf"
            except Exception as format_error:
                print(f"Error generating file in {resume.format} format: {str(format_error)}")
                print("Trying alternative format as fallback...")
                
                # If one format fails, try the other as fallback
                alt_format = "pdf" if resume.format.lower() == "docx" else "docx"
                print(f"Attempting fallback to {alt_format} format...")
                
                try:
                    file_path = resume_generator.create_resume_from_data(resume_data, format=alt_format, template=template)
                    if alt_format == "pdf":
                        media_type = "application/pdf"
                        filename = f"resume_{resume.id}.pdf"
                    else:
                        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        filename = f"resume_{resume.id}.docx"
                    print(f"Successfully generated fallback {alt_format} file")
                except Exception as alt_error:
                    print(f"Fallback generation also failed: {str(alt_error)}")
                    raise ValueError(f"Failed to generate resume in both formats: {str(format_error)}")
            
            if file_path is None:
                print("Failed to generate resume file - file_path is None")
                raise HTTPException(status_code=500, detail="Failed to generate resume file")
                
            print(f"Resume generated successfully at: {file_path}")
            
            # Verify file exists and has content
            if not os.path.exists(file_path):
                print(f"Generated file does not exist at path: {file_path}")
                raise HTTPException(status_code=404, detail="Generated resume file not found")
            
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                print(f"Generated file exists but is empty (0 bytes)")
                raise HTTPException(status_code=500, detail="Generated resume file is empty")
                
            print(f"File verification successful: {file_size} bytes")
                
            # Update file path in database
            print("\nUpdating file path in database...")
            db_resume = db.query(models.Resume).filter(models.Resume.id == resume_id).first()
            if db_resume:
                db_resume.file_path = file_path
                db.commit()
                print("Database updated successfully")
            
            print("\nSending file response...")
            return FileResponse(
                path=file_path,
                media_type=media_type,
                filename=filename
            )
        
        except Exception as e:
            print("\n=== Error in Resume Generation ===")
            print(f"Error: {str(e)}")
            print(f"Error type: {type(e).__name__}")
            import traceback
            print("Full traceback:")
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f"Error generating resume file: {str(e)}")
    
    except HTTPException as he:
        raise he
    except Exception as e:
        print("\n=== Unexpected Error ===")
        print(f"Error: {str(e)}")
        print(f"Error type: {type(e).__name__}")
        import traceback
        print("Full traceback:")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")


@router.get("/history", response_model=List[schemas.Resume])
async def get_resume_history(
    skip: int = 0,
    limit: int = 10,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_active_user)
):
    """Get the resume generation history for the current user"""
    # Use compatible columns
    resumes = db.query(*models.Resume.get_compatible_columns()).filter(
        models.Resume.user_id == current_user.id
    ).order_by(models.Resume.created_at.desc()).offset(skip).limit(limit).all()
    
    # Convert SQLAlchemy result objects to Pydantic models
    return [schemas.Resume(
        id=resume.id,
        user_id=resume.user_id,
        job_analysis_id=resume.job_analysis_id,
        file_path=resume.file_path,
        format=resume.format,
        created_at=resume.created_at
    ) for resume in resumes]


@router.get("/{resume_id}", response_model=schemas.Resume)
async def get_resume(
    resume_id: str,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_active_user)
):
    """Get a specific resume by ID"""
    resume = db.query(*models.Resume.get_compatible_columns()).filter(
        models.Resume.id == resume_id,
        models.Resume.user_id == current_user.id
    ).first()
    
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Convert SQLAlchemy result to Pydantic model
    return schemas.Resume(
        id=resume.id,
        user_id=resume.user_id,
        job_analysis_id=resume.job_analysis_id,
        file_path=resume.file_path,
        format=resume.format,
        created_at=resume.created_at
    )

@router.delete("/{resume_id}", response_model=dict)
async def delete_resume(
    resume_id: str,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_active_user)
):
    """
    Delete a resume by ID.
    """
    # Get the resume from database
    resume = db.query(models.Resume).filter(
        models.Resume.id == resume_id,
        models.Resume.user_id == current_user.id
    ).first()
    
    if not resume:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Resume not found"
        )
    
    try:
        # Delete the physical file if it exists
        if resume.file_path and os.path.exists(resume.file_path):
            os.remove(resume.file_path)
        
        # Delete the database record
        db.delete(resume)
        db.commit()
        
        return {"message": "Resume deleted successfully"}
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error deleting resume: {str(e)}"
        ) 
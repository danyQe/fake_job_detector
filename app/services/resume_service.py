import os
import json
import uuid
from pathlib import Path
from typing import Dict, Any, Optional, List
from sqlalchemy.orm import Session
from app.models import Resume, User
from app.config import settings
from app.services.llm_service import LLMService
import weasyprint
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from bs4 import BeautifulSoup

class ResumeService:
    def __init__(self, db: Session):
        self.db = db
        self.llm_service = LLMService()
        
        # Create resume storage directory if it doesn't exist
        self.resume_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static", "resumes")
        os.makedirs(self.resume_dir, exist_ok=True)
        
        # Define resume templates
        self.templates = {
            "modern": {
                "name": "Modern Professional",
                "description": "Clean and contemporary design with clear section separation",
                "html_template": """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Segoe UI', Arial, sans-serif;
            line-height: 1.6;
            padding: 0.75in;
            color: #333;
            background: white;
            font-size: 10pt;
            max-width: 8.5in;
            margin: 0 auto;
        }}
        
        .header {{
            text-align: center;
            margin-bottom: 22pt;
            padding-bottom: 12pt;
            border-bottom: 1pt solid #eaeaea;
        }}
        
        .name {{
            font-size: 20pt;
            font-weight: bold;
            margin-bottom: 5pt;
            color: #222;
            letter-spacing: 0.5pt;
        }}
        
        .title {{
            font-size: 14pt;
            margin-bottom: 6pt;
            color: #444;
        }}
        
        .contact {{
            font-size: 9.5pt;
            color: #555;
        }}
        
        .section {{
            margin: 18pt 0;
        }}
        
        .section-title {{
            font-size: 13pt;
            font-weight: bold;
            margin-bottom: 10pt;
            border-bottom: 1.5pt solid #444;
            padding-bottom: 3pt;
            color: #222;
            letter-spacing: 0.3pt;
        }}
        
        .company {{
            font-weight: bold;
            margin-bottom: 2pt;
            display: flex;
            justify-content: space-between;
        }}
        
        .company-name {{
            font-weight: bold;
        }}
        
        .job-title {{
            font-weight: bold;
            color: #444;
        }}
        
        .date {{
            font-style: italic;
            font-size: 9.5pt;
            margin-bottom: 5pt;
            color: #555;
        }}
        
        .location {{
            font-size: 9.5pt;
            color: #555;
            margin-bottom: 6pt;
        }}
        
        ul {{
            margin: 8pt 0;
            padding-left: 16pt;
        }}
        
        li {{
            margin-bottom: 5pt;
            list-style-type: disc;
        }}
        
        .skills-list {{
            display: flex;
            flex-wrap: wrap;
            gap: 8pt;
            margin-top: 5pt;
        }}
        
        .skills-list span {{
            background-color: #f2f2f2;
            padding: 3pt 8pt;
            border-radius: 4pt;
            font-size: 9.5pt;
        }}
        
        .job-position {{
            margin-bottom: 14pt;
        }}
        
        .education-item {{
            margin-bottom: 12pt;
        }}
        
        .degree {{
            font-weight: bold;
            margin-bottom: 2pt;
        }}
        
        .school {{
            margin-bottom: 2pt;
        }}
        
        .certification-item {{
            margin-bottom: 10pt;
        }}
        
        .certification-name {{
            font-weight: bold;
            margin-bottom: 2pt;
        }}
        
        @media print {{
            body {{
                padding: 0.5in;
                font-size: 10pt;
                line-height: 1.5;
            }}
            
            .section {{
                page-break-inside: avoid;
            }}
            
            .job-position {{
                page-break-inside: avoid;
            }}
        }}
    </style>
</head>
<body>
    <div class="header">
        <div class="name">{fullName}</div>
        <div class="title">{title}</div>
        <div class="contact">
            {email} | {phone} | {location}{linkedin}
        </div>
    </div>
    
    <div class="section">
        <div class="section-title">Professional Summary</div>
        <p>{summary}</p>
    </div>
    
    <div class="section">
        <div class="section-title">Skills</div>
        <div class="skills-list">
            {skills}
        </div>
    </div>
    
    <div class="section">
        <div class="section-title">Professional Experience</div>
        {experience}
    </div>
    
    <div class="section">
        <div class="section-title">Education</div>
        {education}
    </div>
    
    {certifications}
</body>
</html>"""},
            "classic":
              {
                "name": "Classic Traditional",
                "description": "Timeless and professional layout with emphasis on readability",
                "html_template": """<!DOCTYPE html>
<html>
<head>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Merriweather:wght@400;700&family=Open+Sans:wght@400;600&display=swap');
        
        body {{ 
            font-family: 'Open Sans', sans-serif; 
            line-height: 1.6;
            margin: 0.5in;
            color: #2d3748;
            background-color: #ffffff;
        }}
        .header {{
            text-align: left;
            margin-bottom: 1in;
            padding-bottom: 0.3in;
            border-bottom: 1px solid #e2e8f0;
        }}
        .name {{
            font-family: 'Merriweather', serif;
            font-size: 24pt;
            font-weight: 700;
            color: #1a202c;
            margin-bottom: 0.2in;
            letter-spacing: -0.5px;
        }}
        .title {{
            font-family: 'Merriweather', serif;
            font-size: 14pt;
            font-style: italic;
            color: #4a5568;
            margin-bottom: 0.2in;
        }}
        .contact {{
            font-size: 10pt;
            color: #718096;
            display: flex;
            gap: 1rem;
        }}
        .section {{
            margin-bottom: 0.4in;
        }}
        .section-title {{
            font-family: 'Merriweather', serif;
            font-size: 12pt;
            font-weight: 700;
            color: #2d3748;
            margin-bottom: 0.15in;
            text-transform: uppercase;
            letter-spacing: 0.1em;
            border-bottom: 1px solid #e2e8f0;
            padding-bottom: 0.05in;
        }}
        .company {{
            font-weight: 600;
            color: #2d3748;
            font-size: 11pt;
        }}
        .date {{
            font-style: italic;
            font-size: 9pt;
            color: #718096;
            margin-bottom: 0.1in;
        }}
        ul {{
            margin-top: 0.1in;
            padding-left: 0.3in;
        }}
        li {{
            margin-bottom: 0.1in;
            position: relative;
        }}
        li::before {{
            content: '•';
            color: #4a5568;
            position: absolute;
            left: -0.2in;
        }}
        .skills-list {{
            display: flex;
            flex-wrap: wrap;
            gap: 0.5rem;
        }}
        .skill-tag {{
            background-color: #f7fafc;
            border: 1px solid #e2e8f0;
            padding: 0.2rem 0.5rem;
            border-radius: 0.25rem;
            font-size: 9pt;
            color: #4a5568;
        }}
    </style>
</head>
<body>
    <div class="header">
        <div class="name">{fullName}</div>
        <div class="title">{title}</div>
        <div class="contact">
            <span>{email}</span>
            <span>{phone}</span>
            <span>{location}</span>
            {linkedin}
        </div>
    </div>
    
    <div class="section">
        <div class="section-title">Professional Summary</div>
        <p>{summary}</p>
    </div>
    
    <div class="section">
        <div class="section-title">Skills</div>
        <div class="skills-list">
            {skills}
        </div>
    </div>
    
    <div class="section">
        <div class="section-title">Professional Experience</div>
        {experience}
    </div>
    
    <div class="section">
        <div class="section-title">Education</div>
        {education}
    </div>
    
    {certifications}
</body>
</html>"""
            }
}
    
    # Add new method to handle direct resume_data input
    def create_resume_from_data(self, resume_data: Dict[str, Any], format: str = "pdf", template: str = "modern") -> Optional[str]:
        """
        Generate a resume file from pre-processed resume data.
        
        Args:
            resume_data: Dictionary containing structured resume data
            format: 'pdf' or 'docx'
            template: 'modern' or 'classic' 
            
        Returns:
            Path to the created file
        """
        try:
            # Ensure resume directory exists
            os.makedirs(self.resume_dir, exist_ok=True)
            
            # Generate a filename based on the person's name
            name_part = resume_data.get("fullName", "resume").replace(" ", "_")
            filename = f"{name_part}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            
            print(f"\nGenerating resume file in {format} format with direct data...")
            print(f"Filename: {filename}")
            print(f"Template style: {template}")
            
            # Make sure template is valid
            if template not in ["modern", "classic"]:
                template = "modern"  # Default to modern if invalid
            
            file_path = None
            
            if format.lower() == "pdf":
                file_path = os.path.join(self.resume_dir, f"{filename}.pdf")
                
                # Generate HTML using the selected template
                html_content = self._generate_html_from_data(resume_data, template)
                
                # Create a temporary HTML file
                temp_html = os.path.join(self.resume_dir, f"{filename}_temp.html")
                with open(temp_html, "w", encoding="utf-8") as f:
                    f.write(html_content.strip())
                
                # Generate PDF with WeasyPrint
                html = weasyprint.HTML(filename=temp_html)
                css = weasyprint.CSS(string='''
                    @page {
                        margin: 1cm;
                        size: letter;
                        @top-right { content: "" }
                        @bottom-right { content: "" }
                    }
                ''')
                
                html.write_pdf(
                    file_path,
                    stylesheets=[css],
                    optimize_size=('fonts', 'images'),
                    presentational_hints=True
                )
                
                # Clean up temporary HTML file
                if os.path.exists(temp_html):
                    os.remove(temp_html)
                    
            elif format.lower() == "docx":
                file_path = os.path.join(self.resume_dir, f"{filename}.docx")
                print(f"\nGenerating DOCX file...")
                
                try:
                    # Generate HTML using the selected template if not already generated
                    html_content = self._generate_html_from_data(resume_data, template)
                    
                    # Parse HTML content
                    soup = BeautifulSoup(html_content, 'html.parser')
                    
                    # Create document
                    doc = docx.Document()
                    
                    # Set margins
                    sections = doc.sections
                    for section in sections:
                        section.top_margin = Inches(0.5)
                        section.bottom_margin = Inches(0.5)
                        section.left_margin = Inches(0.5)
                        section.right_margin = Inches(0.5)
                    
                    # Parse the resume content by section
                    # Add the name as a title
                    name_element = soup.find('div', class_='name')
                    if name_element:
                        doc.add_heading(name_element.get_text().strip(), 0)
                    else:
                        doc.add_heading("Resume", 0)
                    
                    # Add job title
                    title_element = soup.find('div', class_='title')
                    if title_element:
                        title_para = doc.add_paragraph()
                        title_run = title_para.add_run(title_element.get_text().strip())
                        title_run.bold = True
                        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # Add contact information
                    contact_element = soup.find('div', class_='contact')
                    if contact_element:
                        contact_para = doc.add_paragraph()
                        contact_para.add_run(contact_element.get_text().strip())
                        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # Add a blank line for spacing
                    doc.add_paragraph()
                    
                    # Process each section in the resume
                    for section in soup.find_all('div', class_='section'):
                        # Add section title
                        section_title = section.find('div', class_='section-title')
                        if section_title:
                            doc.add_heading(section_title.get_text().strip(), 1)
                        
                        # Add summary paragraph if present
                        summary_p = section.find('p')
                        if summary_p and "Summary" in section_title.get_text():
                            doc.add_paragraph(summary_p.get_text().strip())
                        
                        # Add skills list if this is the skills section
                        skills_list = section.find('div', class_='skills-list')
                        if skills_list and "Skills" in section_title.get_text():
                            skills_spans = skills_list.find_all('span')
                            if skills_spans:
                                skills_text = " | ".join([span.get_text().strip() for span in skills_spans])
                                doc.add_paragraph(skills_text)
                        
                        # Add experience items
                        if "Experience" in section_title.get_text() if section_title else False:
                            for job in section.find_all('div', class_='job-position'):
                                # Add company and job title
                                company_div = job.find('div', class_='company')
                                if company_div:
                                    company_name = company_div.find('span', class_='company-name')
                                    date_span = company_div.find('span', class_='date')
                                    
                                    company_para = doc.add_paragraph()
                                    if company_name:
                                        company_run = company_para.add_run(company_name.get_text().strip())
                                        company_run.bold = True
                                    
                                job_title = job.find('div', class_='job-title')
                                if job_title:
                                    title_para = doc.add_paragraph()
                                    title_run = title_para.add_run(job_title.get_text().strip())
                                    title_run.italic = True
                                
                                # Add date if found
                                if date_span:
                                    date_para = doc.add_paragraph()
                                    date_run = date_para.add_run(date_span.get_text().strip())
                                    date_run.italic = True
                                    date_run.font.size = Pt(9.5)
                                
                                # Add job description (bullet points)
                                job_ul = job.find('ul')
                                if job_ul:
                                    for li in job_ul.find_all('li'):
                                        bullet_para = doc.add_paragraph(style='List Bullet')
                                        bullet_para.add_run(li.get_text().strip())
                                else:
                                    # If no bullets, look for paragraph
                                    job_p = job.find('p')
                                    if job_p:
                                        doc.add_paragraph(job_p.get_text().strip())
                        
                        # Add education items
                        if "Education" in section_title.get_text() if section_title else False:
                            for edu in section.find_all('div', class_='education-item'):
                                # Add degree and school
                                degree = edu.find('div', class_='degree')
                                school = edu.find('div', class_='school')
                                date = edu.find('div', class_='date')
                                
                                if school:
                                    school_para = doc.add_paragraph()
                                    school_run = school_para.add_run(school.get_text().strip())
                                    school_run.bold = True
                                
                                if degree:
                                    degree_para = doc.add_paragraph()
                                    degree_para.add_run(degree.get_text().strip())
                                
                                if date:
                                    date_para = doc.add_paragraph()
                                    date_para.add_run(f"Graduated: {date.get_text().strip()}")
                                
                                doc.add_paragraph() # Add space between education entries
                        
                        # Add certifications
                        if "Certification" in section_title.get_text() if section_title else False:
                            for cert in section.find_all('div', class_='certification-item'):
                                cert_name = cert.find('div', class_='certification-name')
                                location = cert.find('div', class_='location')
                                
                                cert_para = doc.add_paragraph()
                                if cert_name and location:
                                    cert_para.add_run(f"{cert_name.get_text().strip()} - {location.get_text().strip()}")
                    
                    # Fallback: process any standalone paragraphs or lists not captured in sections
                    standalone_paras = soup.find_all('p', recursive=False)
                    for para in standalone_paras:
                        doc.add_paragraph(para.get_text().strip())
                    
                    standalone_lists = soup.find_all('ul', recursive=False)
                    for ul in standalone_lists:
                        for li in ul.find_all('li'):
                            doc.add_paragraph(li.get_text().strip(), style='List Bullet')
                            
                    # Save document
                    print(f"Saving DOCX file to: {file_path}")
                    doc.save(file_path)
                    
                    # Verify file
                    if not os.path.exists(file_path):
                        print(f"DOCX file was not created at: {file_path}")
                        raise ValueError("Failed to save DOCX file - file not created")
                    
                    file_size = os.path.getsize(file_path)
                    if file_size == 0:
                        print("DOCX file was created but is empty (0 bytes)")
                        raise ValueError("Generated DOCX file is empty (0 bytes)")
                    
                    print(f"DOCX file created successfully ({file_size} bytes)")
                    
                except Exception as e:
                    print(f"\nError in DOCX generation: {str(e)}")
                    print(f"Error type: {type(e).__name__}")
                    import traceback
                    traceback.print_exc()
                    
                    # Create a very simple fallback DOCX as a last resort
                    try:
                        print("\nAttempting to create simple fallback DOCX...")
                        simple_doc = docx.Document()
                        simple_doc.add_heading("Resume", 0)
                        
                        # Add basic content
                        simple_doc.add_paragraph("This resume was generated as a simple fallback due to formatting errors.")
                        
                        # Add raw text content
                        simple_doc.add_paragraph(soup.get_text())
                        
                        # Save the simple document
                        simple_doc.save(file_path)
                        print("Simple fallback DOCX created")
                    except Exception as fallback_error:
                        print(f"Fallback DOCX also failed: {str(fallback_error)}")
                        raise ValueError(f"Failed to generate DOCX file: {str(e)}")
            else:
                raise ValueError(f"Unsupported format: {format}")
            
            # Verify the file was created successfully
            if file_path is None or not os.path.exists(file_path):
                print(f"Failed to generate file at path: {file_path}")
                raise ValueError(f"Failed to generate resume file in {format} format")
            
            print(f"\nResume generation completed successfully")
            print(f"Output file: {file_path}")
            return file_path
            
        except Exception as e:
            print("\n=== Resume Generation Failed ===")
            print(f"Error: {str(e)}")
            print(f"Error type: {type(e).__name__}")
            import traceback
            print("Full traceback:")
            traceback.print_exc()
            raise ValueError(f"Failed to generate resume: {str(e)}")

    def _generate_html_from_data(self, resume_data: Dict[str, Any], template: str = "modern") -> str:
        """Generate HTML from resume data using the specified template"""
        # Get the template
        template_html = self.templates[template]["html_template"]
        
        # Format skills as spans
        skills_html = ""
        for skill in resume_data.get("skills", []):
            skills_html += f'<span>{skill}</span> '
        
        # Format experience
        experience_html = ""
        for exp in resume_data.get("experience", []):
            current_job = exp.get("current", False)
            end_date = "Present" if current_job else exp.get("endDate", "")
            
            experience_html += f'''
            <div class="job-position">
                <div class="company">
                    <span class="company-name">{exp.get("company", "")}</span>
                    <span class="date">{exp.get("startDate", "")} - {end_date}</span>
                </div>
                <div class="job-title">{exp.get("title", "")}</div>
            '''
            
            # Format description
            description = exp.get("description", "")
            if "\n" in description:
                experience_html += "<ul>"
                for line in description.split("\n"):
                    if line.strip():
                        experience_html += f"<li>{line.strip()}</li>"
                experience_html += "</ul>"
            else:
                experience_html += f"<p>{description}</p>"
            
            experience_html += "</div>"
        
        # Format education
        education_html = ""
        for edu in resume_data.get("education", []):
            education_html += f'''
            <div class="education-item">
                <div class="degree">{edu.get("degree", "")} in {edu.get("field", "")}</div>
                <div class="school">{edu.get("institution", "")}</div>
                <div class="date">{edu.get("graduationDate", "")}</div>
            </div>
            '''
        
        # Format certifications
        certifications_html = ""
        if resume_data.get("certifications"):
            certifications_html = '<div class="section"><div class="section-title">Certifications</div>'
            for cert in resume_data.get("certifications", []):
                certifications_html += f'''
                <div class="certification-item">
                    <div class="certification-name">{cert.get("name", "")}</div>
                    <div class="location">{cert.get("issuer", "")}</div>
                </div>
                '''
            certifications_html += '</div>'
        
        # Format projects
        projects_html = ""
        if resume_data.get("projects"):
            projects_html = '<div class="section"><div class="section-title">Projects</div>'
            for proj in resume_data.get("projects", []):
                projects_html += f'''
                <div class="job-position">
                    <div class="company">
                        <span class="company-name">{proj.get("name", "")}</span>
                        <span class="date">{proj.get("startDate", "")} - {proj.get("endDate", "Present")}</span>
                    </div>
                    <div class="job-title">Technologies: {", ".join(proj.get("technologies", []))}</div>
                '''
                
                # Format description
                description = proj.get("description", "")
                if "\n" in description:
                    projects_html += "<ul>"
                    for line in description.split("\n"):
                        if line.strip():
                            projects_html += f"<li>{line.strip()}</li>"
                    projects_html += "</ul>"
                else:
                    projects_html += f"<p>{description}</p>"
                
                projects_html += "</div>"
            projects_html += '</div>'
        
        # Format LinkedIn
        linkedin_html = ""
        if resume_data.get("linkedin"):
            linkedin_html = f' | {resume_data.get("linkedin", "")}'
        
        # Replace placeholders in template
        formatted_html = template_html.format(
            fullName=resume_data.get("fullName", ""),
            title=resume_data.get("title", ""),
            email=resume_data.get("email", ""),
            phone=resume_data.get("phone", ""),
            location=resume_data.get("location", ""),
            linkedin=linkedin_html,
            summary=resume_data.get("summary", ""),
            skills=skills_html,
            experience=experience_html,
            education=education_html,
            certifications=certifications_html,
            projects=projects_html
        )
        
        return formatted_html
        
    def _apply_modern_docx_styling(self, doc, resume_data: Dict[str, Any]):
        """Apply modern styling to DOCX document"""
        from docx.shared import RGBColor
        
        # Modern style uses a clean, professional look with section separators
        
        # Name and Title
        name = doc.add_paragraph()
        name_run = name.add_run(resume_data.get("fullName", ""))
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_run.font.color.rgb = RGBColor(34, 34, 34)  # Dark gray
        name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        title = doc.add_paragraph()
        title_run = title.add_run(resume_data.get("title", ""))
        title_run.font.size = Pt(14)
        title_run.font.color.rgb = RGBColor(68, 68, 68)  # Medium gray
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Contact Info
        contact = doc.add_paragraph()
        contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        contact_text = f"{resume_data.get('email', '')} | {resume_data.get('phone', '')} | {resume_data.get('location', '')}"
        if resume_data.get("linkedin"):
            contact_text += f" | {resume_data.get('linkedin', '')}"
        contact_run = contact.add_run(contact_text)
        contact_run.font.size = Pt(9.5)
        contact_run.font.color.rgb = RGBColor(85, 85, 85)  # Light gray
        
        # Add section separator - horizontal line
        border_paragraph = doc.add_paragraph()
        border_paragraph.paragraph_format.space_after = Pt(12)
        border_paragraph.paragraph_format.space_before = Pt(12)
        border_run = border_paragraph.add_run()
        border_run.add_break()
        
        # Summary
        self._add_modern_section_heading(doc, "Professional Summary")
        summary_para = doc.add_paragraph()
        summary_para.add_run(resume_data.get("summary", ""))
        
        # Skills
        self._add_modern_section_heading(doc, "Skills")
        skills_para = doc.add_paragraph()
        if resume_data.get("skills"):
            skills_text = " | ".join(resume_data.get("skills", []))
            skills_para.add_run(skills_text)
        
        # Experience
        if resume_data.get("experience"):
            self._add_modern_section_heading(doc, "Professional Experience")
            
            for exp in resume_data.get("experience", []):
                # Company and title
                exp_para = doc.add_paragraph()
                company_run = exp_para.add_run(f"{exp.get('company', '')}")
                company_run.bold = True
                
                # Right-aligned date
                date_text = f"{exp.get('startDate', '')} - "
                date_text += "Present" if exp.get("current") else exp.get('endDate', '')
                
                # Adding title on next line
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"{exp.get('title', '')}")
                title_run.bold = True
                title_run.font.color.rgb = RGBColor(68, 68, 68)  # Medium gray
                
                # Date on separate line
                date_para = doc.add_paragraph()
                date_run = date_para.add_run(date_text)
                date_run.italic = True
                date_run.font.size = Pt(9.5)
                date_run.font.color.rgb = RGBColor(85, 85, 85)  # Light gray
                
                # Description with bullet points
                desc_text = exp.get("description", "")
                if "\n" in desc_text:
                    bullets = desc_text.split("\n")
                    for bullet in bullets:
                        if bullet.strip():
                            bullet_para = doc.add_paragraph(style='List Bullet')
                            bullet_para.add_run(bullet.strip())
                else:
                    doc.add_paragraph(desc_text)
        
        # Education
        if resume_data.get("education"):
            self._add_modern_section_heading(doc, "Education")
            
            for edu in resume_data.get("education", []):
                # Degree
                degree_para = doc.add_paragraph()
                degree_run = degree_para.add_run(f"{edu.get('degree', '')} in {edu.get('field', '')}")
                degree_run.bold = True
                
                # Institution
                school_para = doc.add_paragraph()
                school_run = school_para.add_run(f"{edu.get('institution', '')}")
                
                # Graduation date
                date_para = doc.add_paragraph()
                date_run = date_para.add_run(f"Graduated: {edu.get('graduationDate', '')}")
                date_run.italic = True
                date_run.font.size = Pt(9.5)
                date_run.font.color.rgb = RGBColor(85, 85, 85)  # Light gray
        
        # Certifications
        if resume_data.get("certifications"):
            self._add_modern_section_heading(doc, "Certifications")
            
            for cert in resume_data.get("certifications", []):
                cert_para = doc.add_paragraph()
                cert_run = cert_para.add_run(f"{cert.get('name', '')}")
                cert_run.bold = True
                
                issuer_para = doc.add_paragraph()
                issuer_para.add_run(f"{cert.get('issuer', '')}")
        
        # Projects
        if resume_data.get("projects"):
            self._add_modern_section_heading(doc, "Projects")
            
            for proj in resume_data.get("projects", []):
                name_para = doc.add_paragraph()
                name_run = name_para.add_run(f"{proj.get('name', '')}")
                name_run.bold = True
                
                # Date range
                date_text = f"{proj.get('startDate', '')} - "
                date_text += "Present" if proj.get("current") else proj.get('endDate', '')
                date_para = doc.add_paragraph()
                date_run = date_para.add_run(date_text)
                date_run.italic = True
                date_run.font.size = Pt(9.5)
                
                # Technologies
                if proj.get("technologies"):
                    tech_para = doc.add_paragraph()
                    tech_run = tech_para.add_run(f"Technologies: {', '.join(proj.get('technologies', []))}")
                    tech_run.italic = True
                
                # Description
                desc_text = proj.get("description", "")
                if "\n" in desc_text:
                    bullets = desc_text.split("\n")
                    for bullet in bullets:
                        if bullet.strip():
                            bullet_para = doc.add_paragraph(style='List Bullet')
                            bullet_para.add_run(bullet.strip())
                else:
                    doc.add_paragraph(desc_text)
    
    def _add_modern_section_heading(self, doc, text):
        """Add a modern-styled section heading to a docx document"""
        from docx.shared import RGBColor
        
        heading = doc.add_paragraph()
        heading_run = heading.add_run(text)
        heading_run.bold = True
        heading_run.font.size = Pt(13)
        heading_run.font.color.rgb = RGBColor(34, 34, 34)  # Dark gray
        
        # Add border below
        border_paragraph = doc.add_paragraph()
        border_paragraph.paragraph_format.space_after = Pt(10)
        border_run = border_paragraph.add_run()
        border_paragraph.paragraph_format.bottom_border.width = 1.5
        border_paragraph.paragraph_format.bottom_border.color.rgb = RGBColor(68, 68, 68)
    
    def _apply_classic_docx_styling(self, doc, resume_data: Dict[str, Any]):
        """Apply classic styling to DOCX document"""
        from docx.shared import RGBColor
        
        # Classic style uses a traditional layout with capital letters for headings
        
        # Name and Title
        name = doc.add_paragraph()
        name_run = name.add_run(resume_data.get("fullName", ""))
        name_run.bold = True
        name_run.font.size = Pt(18)
        name.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        title = doc.add_paragraph()
        title_run = title.add_run(resume_data.get("title", ""))
        title_run.font.size = Pt(14)
        title_run.italic = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Contact Info
        contact = doc.add_paragraph()
        contact_text = f"{resume_data.get('email', '')} | {resume_data.get('phone', '')} | {resume_data.get('location', '')}"
        if resume_data.get("linkedin"):
            contact_text += f" | {resume_data.get('linkedin', '')}"
        contact.add_run(contact_text)
        
        doc.add_paragraph()  # Spacer
        
        # Summary
        self._add_classic_section_heading(doc, "PROFESSIONAL SUMMARY")
        doc.add_paragraph(resume_data.get("summary", ""))
        
        # Skills
        self._add_classic_section_heading(doc, "SKILLS")
        skills_para = doc.add_paragraph()
        if resume_data.get("skills"):
            skills_para.add_run(" | ".join(resume_data.get("skills", [])))
        
        # Experience
        if resume_data.get("experience"):
            self._add_classic_section_heading(doc, "PROFESSIONAL EXPERIENCE")
            
            for exp in resume_data.get("experience", []):
                company_para = doc.add_paragraph()
                company_run = company_para.add_run(f"{exp.get('company', '')} - {exp.get('title', '')}")
                company_run.bold = True
                
                date_para = doc.add_paragraph()
                date_text = f"{exp.get('startDate', '')} - "
                date_text += "Present" if exp.get("current") else exp.get('endDate', '')
                date_para.add_run(date_text).italic = True
                
                # Description with bullet points
                desc_text = exp.get("description", "")
                if "\n" in desc_text:
                    bullets = desc_text.split("\n")
                    for bullet in bullets:
                        if bullet.strip():
                            bullet_para = doc.add_paragraph(style='List Bullet')
                            bullet_para.add_run(bullet.strip())
                else:
                    doc.add_paragraph(desc_text)
        
        # Education
        if resume_data.get("education"):
            self._add_classic_section_heading(doc, "EDUCATION")
            
            for edu in resume_data.get("education", []):
                edu_para = doc.add_paragraph()
                edu_run = edu_para.add_run(f"{edu.get('institution', '')} - {edu.get('degree', '')} in {edu.get('field', '')}")
                edu_run.bold = True
                
                grad_para = doc.add_paragraph()
                grad_para.add_run(f"Graduated: {edu.get('graduationDate', '')}").italic = True
        
        # Certifications
        if resume_data.get("certifications"):
            self._add_classic_section_heading(doc, "CERTIFICATIONS")
            
            for cert in resume_data.get("certifications", []):
                cert_para = doc.add_paragraph(style='List Bullet')
                cert_para.add_run(f"{cert.get('name', '')} - {cert.get('issuer', '')}")
        
        # Projects
        if resume_data.get("projects"):
            self._add_classic_section_heading(doc, "PROJECTS")
            
            for proj in resume_data.get("projects", []):
                name_para = doc.add_paragraph()
                name_run = name_para.add_run(f"{proj.get('name', '')}")
                name_run.bold = True
                
                # Date range
                date_para = doc.add_paragraph()
                date_text = f"{proj.get('startDate', '')} - "
                date_text += "Present" if proj.get("current") else proj.get('endDate', '')
                date_para.add_run(date_text).italic = True
                
                # Technologies
                if proj.get("technologies"):
                    tech_para = doc.add_paragraph()
                    tech_para.add_run(f"Technologies: {', '.join(proj.get('technologies', []))}")
                
                # Description
                desc_text = proj.get("description", "")
                if "\n" in desc_text:
                    bullets = desc_text.split("\n")
                    for bullet in bullets:
                        if bullet.strip():
                            bullet_para = doc.add_paragraph(style='List Bullet')
                            bullet_para.add_run(bullet.strip())
                else:
                    doc.add_paragraph(desc_text)
    
    def _add_classic_section_heading(self, doc, text):
        """Add a classic-styled section heading to a docx document"""
        heading = doc.add_paragraph()
        heading_run = heading.add_run(text)
        heading_run.bold = True
        heading_run.font.small_caps = True
        
        # Underline the heading
        heading_run.underline = True
    
    async def generate_resume(
        self,
        user_id: str,
        job_details: Dict[str, Any],
        personal_info: Dict[str, Any],
        format: str = "pdf",
        template: str = "modern"
    ) -> Optional[str]:
        """
        Generate a resume based on job details and personal information.
        Returns the resume ID if successful, None otherwise.
        """
        try:
            # Create prompt for LLM
            prompt = self._create_resume_prompt(job_details, personal_info)
            
            # Generate resume content with LLM
            resume_content = await self.llm_service.generate_text(prompt)
            
            if not resume_content:
                return None
            
            # Create resume file (PDF or DOCX)
            file_path = self._create_resume_file(resume_content, personal_info["fullName"], format, template)
            
            if not file_path:
                return None
            
            # Create resume record in database
            resume = Resume(
                user_id=user_id,
                job_analysis_id=job_details.get("id"),
                file_path=file_path,
                format=format
            )
            
            self.db.add(resume)
            self.db.commit()
            self.db.refresh(resume)
            
            return resume.id
        
        except Exception as e:
            print(f"Error generating resume: {str(e)}")
            return None
    
    def _create_resume_prompt(self, job_details: Dict[str, Any], personal_info: Dict[str, Any]) -> str:
        """
        Create a prompt for the LLM to generate a resume.
        """
        prompt = f"""
        Create a professional resume for {personal_info['fullName']} tailored to this job:
        
        JOB TITLE: {job_details['title']}
        COMPANY: {job_details['company']}
        JOB DESCRIPTION: {job_details['description']}
        
        JOB ANALYSIS: {job_details.get('analysis', 'No analysis available')}
        
        CANDIDATE INFORMATION:
        - Full Name: {personal_info['fullName']}
        - Professional Title: {personal_info['title']}
        - Email: {personal_info['email']}
        - Phone: {personal_info['phone']}
        - Location: {personal_info['location']}
        - LinkedIn: {personal_info.get('linkedin', 'Not provided')}
        
        PROFESSIONAL SUMMARY:
        {personal_info['summary']}
        
        SKILLS:
        {', '.join(personal_info['skills'])}
        
        EXPERIENCE:
        """
        
        for exp in personal_info['experience']:
            current = "Present" if exp.get('current', False) else exp.get('endDate', 'N/A')
            prompt += f"""
            - {exp['title']} at {exp['company']} ({exp['startDate']} - {current})
              {exp['description']}
            """
        
        prompt += "\nEDUCATION:\n"
        for edu in personal_info['education']:
            prompt += f"""
            - {edu['degree']} in {edu.get('field', 'N/A')} from {edu['institution']} ({edu['graduationDate']})
            """
        
        if personal_info.get('certifications'):
            prompt += "\nCERTIFICATIONS:\n"
            for cert in personal_info['certifications']:
                prompt += f"""
                - {cert['name']} from {cert['issuer']}
                """
        
        prompt += """
        INSTRUCTIONS:
        1. Format the resume professionally, with clear section headings and consistent spacing.
        2. Highlight skills and experiences that match the job requirements.
        3. Use bullet points for clarity and conciseness.
        4. Quantify achievements with metrics when possible.
        5. Ensure there are no spelling or grammar errors.
        6. The resume should be ATS-friendly.
        7. Return the resume in a format ready to be converted to a PDF or DOCX file.
        """
        
        return prompt
    
    def _create_resume_file(self, content: str, name: str, format: str, template: str = "modern") -> Optional[str]:
        """
        Create a resume file in the specified format.
        
        Args:
            content: HTML content for the resume
            name: Base filename
            format: 'pdf' or 'docx'
            template: 'modern' or 'classic' template to use (optional)
            
        Returns:
            Path to the created file
        """
        os.makedirs(self.resume_dir, exist_ok=True)
        filename = f"{name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        try:
            print(f"\nCreating resume file in {format} format...")
            print(f"Filename: {filename}")
            
            # Store content in html_content for consistency
            html_content = content
            
            if format.lower() == "pdf":
                try:
                    file_path = os.path.join(self.resume_dir, f"{filename}.pdf")
                    temp_html = os.path.join(self.resume_dir, f"{filename}_temp.html")
                    
                    try:
                        print("\nStarting PDF generation...")
                        print(f"Creating temporary HTML file: {temp_html}")
                        
                        # Create a temporary HTML file with proper encoding
                        with open(temp_html, "w", encoding="utf-8") as f:
                            f.write(html_content.strip())
                        print("Temporary HTML file created successfully")
                        
                        print("\nGenerating PDF with WeasyPrint...")
                        # Generate PDF using WeasyPrint with explicit configuration
                        html = weasyprint.HTML(filename=temp_html)
                        print("HTML loaded in WeasyPrint")
                        
                        css = weasyprint.CSS(string='''
                            @page {
                                margin: 1cm;
                                size: letter;
                                @top-right { content: "" }
                                @bottom-right { content: "" }
                            }
                        ''')
                        print("CSS configuration created")
                        
                        # Write PDF with custom CSS
                        print(f"Writing PDF to: {file_path}")
                        html.write_pdf(
                            file_path,
                            stylesheets=[css],
                            optimize_size=('fonts', 'images'),
                            presentational_hints=True
                        )
                        print("PDF generated successfully")
                        
                    except Exception as e:
                        print(f"\nError in PDF generation: {str(e)}")
                        print(f"Error type: {type(e).__name__}")
                        print("Falling back to DOCX format...")
                        
                        # Fall back to DOCX format
                        format = "docx"
                        file_path = os.path.join(self.resume_dir, f"{filename}.docx")
                        print("\nGenerating DOCX file as fallback...")
                        doc = docx.Document()
                        
                        # Add basic content from HTML
                        doc.add_heading("Resume", 0)
                        doc.add_paragraph(f"Note: This is a fallback DOCX version due to PDF generation error: {str(e)}")
                        doc.add_paragraph("Please install GTK for Windows to enable PDF generation.")
                        
                        # Add some content from the HTML
                        soup = BeautifulSoup(html_content, 'html.parser')
                        
                        # Extract text content from HTML and add to document
                        for para in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                            if para.name.startswith('h'):
                                doc.add_heading(para.get_text(), level=int(para.name[1]))
                            else:
                                doc.add_paragraph(para.get_text())
                        
                        doc.save(file_path)
                        print("DOCX fallback file generated successfully")
                        
                    finally:
                        # Clean up temporary HTML file
                        if os.path.exists(temp_html):
                            try:
                                os.remove(temp_html)
                                print("Temporary HTML file cleaned up")
                            except Exception as e:
                                print(f"Error cleaning up temporary file: {str(e)}")
                except Exception as e:
                    print(f"Error in PDF/fallback generation: {str(e)}")
                    raise
            
            if format.lower() == "docx":
                file_path = os.path.join(self.resume_dir, f"{filename}.docx")
                print(f"\nGenerating DOCX file...")
                
                try:
                    # Generate HTML using the selected template if not already generated
                    html_content = self._generate_html_from_data(resume_data, template)
                    
                    # Parse HTML content
                    soup = BeautifulSoup(html_content, 'html.parser')
                    
                    # Create document
                    doc = docx.Document()
                    
                    # Set margins
                    sections = doc.sections
                    for section in sections:
                        section.top_margin = Inches(0.5)
                        section.bottom_margin = Inches(0.5)
                        section.left_margin = Inches(0.5)
                        section.right_margin = Inches(0.5)
                    
                    # Parse the resume content by section
                    # Add the name as a title
                    name_element = soup.find('div', class_='name')
                    if name_element:
                        doc.add_heading(name_element.get_text().strip(), 0)
                    else:
                        doc.add_heading("Resume", 0)
                    
                    # Add job title
                    title_element = soup.find('div', class_='title')
                    if title_element:
                        title_para = doc.add_paragraph()
                        title_run = title_para.add_run(title_element.get_text().strip())
                        title_run.bold = True
                        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # Add contact information
                    contact_element = soup.find('div', class_='contact')
                    if contact_element:
                        contact_para = doc.add_paragraph()
                        contact_para.add_run(contact_element.get_text().strip())
                        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # Add a blank line for spacing
                    doc.add_paragraph()
                    
                    # Process each section in the resume
                    for section in soup.find_all('div', class_='section'):
                        # Add section title
                        section_title = section.find('div', class_='section-title')
                        if section_title:
                            doc.add_heading(section_title.get_text().strip(), 1)
                        
                        # Add summary paragraph if present
                        summary_p = section.find('p')
                        if summary_p and "Summary" in section_title.get_text():
                            doc.add_paragraph(summary_p.get_text().strip())
                        
                        # Add skills list if this is the skills section
                        skills_list = section.find('div', class_='skills-list')
                        if skills_list and "Skills" in section_title.get_text():
                            skills_spans = skills_list.find_all('span')
                            if skills_spans:
                                skills_text = " | ".join([span.get_text().strip() for span in skills_spans])
                                doc.add_paragraph(skills_text)
                        
                        # Add experience items
                        if "Experience" in section_title.get_text() if section_title else False:
                            for job in section.find_all('div', class_='job-position'):
                                # Add company and job title
                                company_div = job.find('div', class_='company')
                                if company_div:
                                    company_name = company_div.find('span', class_='company-name')
                                    date_span = company_div.find('span', class_='date')
                                    
                                    company_para = doc.add_paragraph()
                                    if company_name:
                                        company_run = company_para.add_run(company_name.get_text().strip())
                                        company_run.bold = True
                                    
                                job_title = job.find('div', class_='job-title')
                                if job_title:
                                    title_para = doc.add_paragraph()
                                    title_run = title_para.add_run(job_title.get_text().strip())
                                    title_run.italic = True
                                
                                # Add date if found
                                if date_span:
                                    date_para = doc.add_paragraph()
                                    date_run = date_para.add_run(date_span.get_text().strip())
                                    date_run.italic = True
                                    date_run.font.size = Pt(9.5)
                                
                                # Add job description (bullet points)
                                job_ul = job.find('ul')
                                if job_ul:
                                    for li in job_ul.find_all('li'):
                                        bullet_para = doc.add_paragraph(style='List Bullet')
                                        bullet_para.add_run(li.get_text().strip())
                                else:
                                    # If no bullets, look for paragraph
                                    job_p = job.find('p')
                                    if job_p:
                                        doc.add_paragraph(job_p.get_text().strip())
                        
                        # Add education items
                        if "Education" in section_title.get_text() if section_title else False:
                            for edu in section.find_all('div', class_='education-item'):
                                # Add degree and school
                                degree = edu.find('div', class_='degree')
                                school = edu.find('div', class_='school')
                                date = edu.find('div', class_='date')
                                
                                if school:
                                    school_para = doc.add_paragraph()
                                    school_run = school_para.add_run(school.get_text().strip())
                                    school_run.bold = True
                                
                                if degree:
                                    degree_para = doc.add_paragraph()
                                    degree_para.add_run(degree.get_text().strip())
                                
                                if date:
                                    date_para = doc.add_paragraph()
                                    date_para.add_run(f"Graduated: {date.get_text().strip()}")
                                
                                doc.add_paragraph() # Add space between education entries
                        
                        # Add certifications
                        if "Certification" in section_title.get_text() if section_title else False:
                            for cert in section.find_all('div', class_='certification-item'):
                                cert_name = cert.find('div', class_='certification-name')
                                location = cert.find('div', class_='location')
                                
                                cert_para = doc.add_paragraph()
                                if cert_name and location:
                                    cert_para.add_run(f"{cert_name.get_text().strip()} - {location.get_text().strip()}")
                    
                    # Fallback: process any standalone paragraphs or lists not captured in sections
                    standalone_paras = soup.find_all('p', recursive=False)
                    for para in standalone_paras:
                        doc.add_paragraph(para.get_text().strip())
                    
                    standalone_lists = soup.find_all('ul', recursive=False)
                    for ul in standalone_lists:
                        for li in ul.find_all('li'):
                            doc.add_paragraph(li.get_text().strip(), style='List Bullet')
                            
                    # Save document
                    print(f"Saving DOCX file to: {file_path}")
                    doc.save(file_path)
                    
                    # Verify file
                    if not os.path.exists(file_path):
                        print(f"DOCX file was not created at: {file_path}")
                        raise ValueError("Failed to save DOCX file - file not created")
                    
                    file_size = os.path.getsize(file_path)
                    if file_size == 0:
                        print("DOCX file was created but is empty (0 bytes)")
                        raise ValueError("Generated DOCX file is empty (0 bytes)")
                    
                    print(f"DOCX file created successfully ({file_size} bytes)")
                    
                except Exception as e:
                    print(f"\nError in DOCX generation: {str(e)}")
                    print(f"Error type: {type(e).__name__}")
                    import traceback
                    traceback.print_exc()
                    
                    # Create a very simple fallback DOCX as a last resort
                    try:
                        print("\nAttempting to create simple fallback DOCX...")
                        simple_doc = docx.Document()
                        simple_doc.add_heading("Resume", 0)
                        
                        # Add basic content
                        simple_doc.add_paragraph("This resume was generated as a simple fallback due to formatting errors.")
                        
                        # Add raw text content
                        simple_doc.add_paragraph(soup.get_text())
                        
                        # Save the simple document
                        simple_doc.save(file_path)
                        print("Simple fallback DOCX created")
                    except Exception as fallback_error:
                        print(f"Fallback DOCX also failed: {str(fallback_error)}")
                        raise ValueError(f"Failed to generate DOCX file: {str(e)}")
            else:
                raise ValueError(f"Unsupported format: {format}")
            
            print(f"\nResume generation completed successfully")
            print(f"Output file: {file_path}")
            
            # Final verification
            if not os.path.exists(file_path):
                raise ValueError(f"File was not created at: {file_path}")
            
            return file_path
            
        except Exception as e:
            print("\n=== Resume Generation Failed ===")
            print(f"Error: {str(e)}")
            print(f"Error type: {type(e).__name__}")
            import traceback
            print("Full traceback:")
            traceback.print_exc()
            raise 